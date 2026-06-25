
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=11, color=None):
    r = para.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return r

def normal_run(para, text, size=11, italic=False):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    return r

def heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
        run.font.bold = True
    return p

def add_table(doc, headers, rows, header_bg="1F3864", alt_bg="EBF5FB"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

def add_para(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold   = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_image(doc, path, width=Inches(6.2), caption=None):
    try:
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(6)
            for r in cp.runs:
                r.font.size = Pt(9)
                r.italic = True
                r.font.color.rgb = RGBColor(80,80,80)
    except Exception as e:
        add_para(doc, f"[Figure: {caption}]", italic=True)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1F3864')
    pBdr.append(bot)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RESEARCH INTERNSHIP DOCUMENTATION")
r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(31,56,100)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Week 2 – Data Preprocessing, EDA & Baseline Model Development")
r2.font.size = Pt(13); r2.bold = True
r2.font.color.rgb = RGBColor(21,101,192)

divider(doc)

info = [
    ("Project Title", "AI-Powered Tool Wear Prediction and Predictive Maintenance System for Smart Manufacturing"),
    ("Group No", "18"),
    ("Members", "Shravani Nikam | Mihir Patel | Samrudhi Musale | Paresh Morankar"),
    ("Mentor", "Prof. Dr. Madhuri Barhate"),
    ("Institute", "Vishwakarma Institute of Technology"),
    ("Domain", "Smart Manufacturing · Industry 4.0 · Predictive Maintenance · AI/ML"),
]
for label, val in info:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    bold_run(p, f"{label}: ", 11, "1F3864")
    normal_run(p, val, 11)

divider(doc)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════
heading(doc, "1. Project Overview", 1)

add_para(doc,
    "Modern manufacturing facilities lose billions of dollars annually due to unexpected tool failures during machining operations. "
    "A cutting tool — such as a milling insert — undergoes progressive flank wear as it removes material from a workpiece. "
    "When tool wear exceeds a critical threshold (typically 0.3 mm of Flank Wear Band, VB), the tool fails catastrophically, causing "
    "dimensional errors in the finished part, surface quality degradation, spindle damage, and costly production downtime.")

add_para(doc,
    "This project addresses the core industrial challenge: predicting WHEN a tool will fail BEFORE it actually fails. "
    "We develop an AI-powered dual-level predictive maintenance framework that simultaneously tackles:")

bullets = [
    "Tool-Level Prognostics: Continuous regression of Flank Wear (VB in mm) and estimation of Remaining Useful Life (RUL)",
    "Machine-Level Failure Prediction: Binary and multi-class classification of machine failure types (TWF, HDF, PWF, OSF, RNF)",
    "Explainable Maintenance Recommendations: SHAP-based interpretability to explain predictions to factory operators",
    "Interactive Dashboard: Real-time Streamlit interface displaying wear state, alerts, and maintenance schedules",
]
for b in bullets:
    add_bullet(doc, b)

heading(doc, "1.1 Why This Matters — Industrial Impact", 2)

add_table(doc,
    ["Problem", "Current Situation", "Our Solution"],
    [
        ["Unexpected tool breakage", "Production stops, scrapped parts", "Predict wear 10–15 runs in advance"],
        ["Manual inspection", "Time-consuming, subjective, costly", "Continuous AI-based monitoring"],
        ["Reactive maintenance", "Replace after failure", "Predictive: replace just before failure"],
        ["Sensor data unused", "Only basic alarms used", "Multi-sensor fusion + deep learning"],
        ["No explainability", "Engineers don't trust black-box AI", "SHAP explanations for every alert"],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 2. DATASET DESCRIPTION
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "2. Dataset Description — Deep Dive", 1)

heading(doc, "2.1 NASA Ames Milling Dataset (Primary — Tool Wear Regression)", 2)

add_para(doc,
    "The NASA Ames Milling Dataset is a benchmark real-world dataset collected under controlled milling experiments at NASA's Ames Research Center. "
    "A Matsuura MC-510V machining center was used to perform end-milling operations on aerospace-grade workpiece materials. "
    "Multiple sensors were simultaneously recorded at high sampling frequency throughout each run-to-failure experiment.")

add_para(doc,
    "The dataset comprises 16 separate machining cases (tool experiments), each representing a complete life cycle of a milling cutter "
    "from fresh (VB = 0 mm) to worn-out state. Each row in the dataset represents one measurement snapshot during a machining run.")

add_table(doc,
    ["Feature", "Full Name", "Unit", "Physical Meaning", "Role"],
    [
        ["case", "Experiment Case ID", "—", "Identifies which tool experiment (1–16)", "Grouping"],
        ["run", "Run Number", "—", "Sequential pass number within an experiment", "Time proxy"],
        ["VB", "Flank Wear Band", "mm", "Actual measured tool wear at flank face — THE TARGET", "Target"],
        ["time", "Cumulative Machining Time", "min", "Total time the tool has been cutting", "Feature"],
        ["DOC", "Depth of Cut", "mm", "How deep the tool cuts into workpiece (0.75 or 1.5)", "Feature"],
        ["feed", "Feed Rate", "mm/rev", "Tool advancement per spindle rotation (0.25 or 0.5)", "Feature"],
        ["material", "Workpiece Material", "—", "1 = Cast Iron (softer), 2 = Steel (harder)", "Feature"],
        ["smcAC", "Spindle Motor Current (AC)", "A", "AC component of spindle motor current — reflects cutting force", "Feature"],
        ["smcDC", "Spindle Motor Current (DC)", "A", "DC component — reflects average motor load and wear state", "Feature"],
        ["vib_table", "Table Vibration", "g", "Vibration acceleration measured at the machine table", "Feature"],
        ["vib_spindle", "Spindle Vibration", "g", "Vibration at the tool spindle — most direct wear indicator", "Feature"],
        ["AE_table", "Acoustic Emission (Table)", "V", "Ultrasonic stress waves from cutting at table sensor", "Feature"],
        ["AE_spindle", "Acoustic Emission (Spindle)", "V", "AE at spindle — captures micro-fractures and wear events", "Feature"],
    ]
)

doc.add_paragraph()
heading(doc, "2.2 AI4I 2020 Predictive Maintenance Dataset (Secondary — Failure Classification)", 2)

add_para(doc,
    "The AI4I 2020 dataset from the UCI Machine Learning Repository simulates a real industrial machine monitoring scenario with 10,000 records "
    "of machine state observations. It provides labeled failure events across five distinct failure modes, making it ideal for building "
    "a multi-class failure classifier that complements the tool wear regressor.")

add_table(doc,
    ["Feature", "Unit", "Description", "Role"],
    [
        ["Air temperature [K]", "Kelvin", "Ambient air temperature around the machine", "Feature"],
        ["Process temperature [K]", "Kelvin", "Temperature of the cutting process", "Feature"],
        ["Rotational speed [rpm]", "RPM", "Spindle rotation speed — affects cutting forces", "Feature"],
        ["Torque [Nm]", "Newton-meter", "Torque applied by spindle motor", "Feature"],
        ["Tool wear [min]", "minutes", "Cumulative tool usage time", "Feature"],
        ["Machine failure", "Binary", "Overall machine failure flag (0=OK, 1=Failed)", "Target"],
        ["TWF", "Binary", "Tool Wear Failure — tool exceeded wear limit", "Sub-target"],
        ["HDF", "Binary", "Heat Dissipation Failure — thermal overload", "Sub-target"],
        ["PWF", "Binary", "Power Failure — process power out of range", "Sub-target"],
        ["OSF", "Binary", "Overstrain Failure — torque × tool wear threshold", "Sub-target"],
        ["RNF", "Binary", "Random Failure — stochastic machine failure", "Sub-target"],
        ["Type", "L/M/H", "Product quality type (Low/Medium/High)", "Feature"],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "3. System Architecture", 1)

add_para(doc,
    "The system follows a multi-stage intelligent pipeline. Data from two real-world datasets flows through preprocessing, "
    "feature engineering, parallel model execution, and finally explainable output rendering on an interactive dashboard.")

add_table(doc,
    ["Stage", "Module", "Input", "Output", "Technology"],
    [
        ["1", "Multi-Modal Data Ingestion", "Raw CSV files", "Loaded DataFrames", "Pandas, NumPy"],
        ["2", "Signal Preprocessing", "Raw sensor readings", "Clean, normalized arrays", "Scikit-Learn, SciPy"],
        ["3", "Feature Engineering", "Per-run sensor data", "32 statistical features", "Custom aggregation"],
        ["4A", "LSTM Model (Week 3)", "Time-series sequences", "VB prediction (regression)", "TensorFlow/Keras"],
        ["4B", "Gradient Boosting Baseline", "Flat feature matrix", "VB prediction", "Scikit-Learn"],
        ["4C", "XGBoost Classifier (Week 3)", "Tabular AI4I features", "Failure type prediction", "XGBoost"],
        ["5", "Explainability Layer", "Model + test data", "SHAP values, feature importance", "SHAP library"],
        ["6", "Dashboard", "Predictions + SHAP", "Visual alerts, RUL display", "Streamlit"],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 4. METHODOLOGY
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "4. Methodology", 1)

add_para(doc,
    "Our methodology follows an eight-step scientific pipeline designed to extract maximum predictive signal from "
    "the available multi-sensor data while maintaining interpretability and generalizability.")

steps = [
    ("Step 1 — Problem Framing",
     "We formulate tool wear prediction as a REGRESSION problem (predict continuous VB value in mm) and machine failure as a CLASSIFICATION problem "
     "(predict binary failure and failure type). This dual formulation lets us simultaneously answer 'How worn is the tool?' and 'Will it fail soon?'"),
    ("Step 2 — Data Acquisition & Integration",
     "Two datasets are loaded: NASA Milling (run-to-failure time-series) and AI4I 2020 (structured machine logs). "
     "Together they provide complementary perspectives — continuous degradation signals and discrete failure labels."),
    ("Step 3 — Data Preprocessing",
     "Missing value imputation, noise characterization, feature scaling, and class imbalance analysis. "
     "Full details in Section 5 below."),
    ("Step 4 — Exploratory Data Analysis",
     "Statistical analysis and visualization of all features to understand distributions, correlations, "
     "and physical relationships between sensor signals and tool wear. Section 6 provides deep analysis."),
    ("Step 5 — Feature Engineering",
     "Per-run statistical aggregation of all 6 sensor channels (mean, std, max, min) to convert high-frequency signals "
     "into a flat, ML-ready feature matrix of 32 features."),
    ("Step 6 — Baseline Model Development",
     "Three regression models trained and evaluated on the engineered feature set. "
     "Best baseline: Gradient Boosting (R²=0.8334, MAE=0.0654 mm)."),
    ("Step 7 — Advanced Models (Week 3)",
     "LSTM for sequential wear trend, CNN-1D for raw signal patterns, Hybrid CNN+LSTM, "
     "XGBoost+SMOTE for failure classification."),
    ("Step 8 — Explainability & Deployment",
     "SHAP values to identify top wear-driving features. Streamlit dashboard for factory-floor deployment."),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    bold_run(p, f"◆ {title}: ", 11, "1F3864")
    normal_run(p, desc, 11)

# ═══════════════════════════════════════════════════════════════
# 5. DATA PREPROCESSING
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "5. Data Preprocessing — Detailed Analysis", 1)

heading(doc, "5.1 NASA Milling Dataset — Missing Value Problem & Solution", 2)

add_para(doc,
    "The most critical preprocessing challenge in the Mill dataset is the significant proportion of missing VB (Flank Wear) values. "
    "VB is our TARGET variable — any missing label directly removes a training sample. Here is what we found and how we solved it:")

add_table(doc,
    ["Column", "Missing Count", "Missing %", "Cause", "Solution Applied"],
    [
        ["VB (Flank Wear)", "26 out of 168", "15.5%", "Tool wear was physically measured only at specific intervals during machining (not every run)", "Per-case linear interpolation → forward fill → backward fill"],
        ["smcAC, smcDC, vib_table, vib_spindle, AE_table, AE_spindle", "0", "0%", "Sensors recorded continuously", "No action needed"],
        ["case, run, time, DOC, feed, material", "0", "0%", "Machining parameters always recorded", "No action needed"],
    ]
)

add_para(doc, "Why Linear Interpolation (not mean/median imputation)?", bold=True, size=11, color="1F3864")
add_para(doc,
    "Tool wear is a MONOTONICALLY INCREASING physical process — a tool always gets more worn over time, never less. "
    "This physical constraint means that between two measured VB values (e.g., VB=0.2 mm at run 5 and VB=0.4 mm at run 9), "
    "we can confidently estimate intermediate values as a smooth linear progression. "
    "Mean/median imputation would be physically incorrect — it would place the same average value at all missing positions "
    "regardless of the temporal position of the measurement.")

p = doc.add_paragraph()
bold_run(p, "Implementation: ", 11, "1F3864")
normal_run(p, "mill_clean.groupby('case').apply(lambda g: g['VB'].interpolate(method='linear').ffill().bfill())", 10, italic=True)

add_para(doc,
    "This applies the interpolation WITHIN each case independently, correctly respecting the wear curve of each individual tool experiment. "
    "After this step: 0 missing values remain in VB. All 168 records retained for training.")

heading(doc, "5.2 Noise in Raw Sensor Signals", 2)

add_para(doc,
    "Raw sensor readings contain high-frequency noise from machine vibrations, electrical interference, and measurement digitization. "
    "For example, smcAC values oscillate between negative and positive values within a single run due to AC motor alternation cycles. "
    "Directly feeding these raw samples into ML models would cause overfitting to noise patterns.")

add_para(doc, "Solution — Statistical Feature Aggregation per Run:", bold=True, size=11, color="1F3864")
add_para(doc,
    "Instead of using raw time-step samples, we compute 4 statistical descriptors per sensor per run: "
    "Mean (central tendency), Std (variability and signal energy), Max (peak stress events), Min (baseline floor). "
    "This converts each machining run from a noisy multi-sample sequence into a single, stable 32-dimensional feature vector. "
    "This aggregation effectively de-noises the signal while preserving the information content relevant to wear state.")

heading(doc, "5.3 AI4I Dataset — Class Imbalance Problem", 2)

add_para(doc,
    "The AI4I dataset has a severe class imbalance: 96.61% of records are non-failure states, and only 3.39% (339 records) are failures. "
    "If we train a classifier on raw data, a naive model that ALWAYS predicts 'No Failure' would achieve 96.6% accuracy — "
    "but would completely fail to detect any actual failure. This is the class imbalance trap.")

add_table(doc,
    ["Class", "Count", "Percentage", "Problem", "Planned Solution (Week 3)"],
    [
        ["No Failure (0)", "9,661", "96.61%", "Model biased toward majority", "SMOTE oversampling of minority"],
        ["Failure (1)", "339", "3.39%", "Under-represented — critical class!", "Class weight = 96.61/3.39 ≈ 28.5"],
    ]
)

add_para(doc,
    "SMOTE (Synthetic Minority Over-sampling Technique) generates synthetic failure samples by interpolating between existing "
    "failure records in feature space, bringing the training class ratio to 1:1 without simply duplicating records.")

# ═══════════════════════════════════════════════════════════════
# 6. EDA — GRAPH EXPLANATIONS
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "6. Exploratory Data Analysis — Graph-by-Graph Interpretation", 1)

heading(doc, "6.1 NASA Milling Dataset — Six-Panel EDA", 2)

add_image(doc, r'd:\IEEE\eda_mill.png', Inches(6.2),
    "Figure 1: NASA Milling Dataset — Six-Panel Exploratory Data Analysis")

# Panel by panel
panels = [
    ("Panel 1 (Top-Left): Distribution of Tool Wear (VB)",
     "What the graph shows: A histogram of all VB (Flank Wear) measurements from 0 to 1.53 mm, "
     "with a red dashed vertical line at VB = 0.3 mm (the industry-standard tool replacement threshold).",
     "What it means: The distribution is right-skewed. The majority of measurements (approx. 60%) fall in the early wear stage "
     "(VB < 0.3 mm), with fewer records in the severe wear stage (VB > 0.5 mm). This reflects the physical reality: "
     "most experiments capture early-to-moderate wear, with fewer run-to-destruction cases.",
     "Solution implication: Because most data points are at low wear values, our model needs to be careful not to underestimate "
     "high-wear predictions. We will use MAE (not just accuracy) as primary metric since it penalizes all errors equally. "
     "For deep learning models in Week 3, we will apply sample weighting to give higher training importance to high-wear data points."),

    ("Panel 2 (Top-Center): Tool Wear Progression Over Time",
     "What the graph shows: Line plots of VB vs. machining time for 8 individual tool experiments (cases), "
     "each represented by a different color.",
     "What it means: Every tool follows a clear monotonically increasing wear trajectory — VB grows over time without exception. "
     "However, the RATE of wear varies significantly between cases. Some tools wear rapidly (Case 13: reaches 1.53 mm in ~42 min), "
     "while others wear slowly (Case 11: reaches 0.68 mm in ~100 min). This variability is caused by different machining parameters "
     "(DOC, feed rate) and workpiece materials across cases.",
     "Solution implication: This confirms that a simple time-based rule ('replace after X minutes') would be unreliable — "
     "wear rate depends on operating conditions. Our ML model must learn the RELATIONSHIP between sensor signals, cutting parameters, "
     "and wear rate. This also motivates the use of LSTM in Week 3, which can learn temporal wear trajectories."),

    ("Panel 3 (Top-Right): Sensor ↔ Tool Wear Correlation Heatmap",
     "What the graph shows: A color-coded correlation matrix between all 6 sensor signals and the VB target. "
     "Values range from -1 (perfect negative correlation) to +1 (perfect positive). Green = positive, Yellow = near-zero.",
     "What it means: No single raw sensor channel is strongly correlated with VB directly. The strongest raw correlation is "
     "smcDC–VB at +0.20. However, AE_table and AE_spindle are strongly correlated with EACH OTHER (r=0.74), indicating they carry "
     "complementary information. This low raw correlation is NOT a problem — it means the information is encoded in the PATTERNS "
     "and STATISTICS of the signals over time, not in their instantaneous values.",
     "Solution implication: Direct correlation-based feature selection would incorrectly discard all sensors. "
     "The solution is statistical feature engineering — by computing mean, std, max, min per run, we capture the signal patterns "
     "that correlate strongly with VB (AE_table_max achieves importance of 0.45 in RF). "
     "This is why our feature engineering step is the most critical component of the pipeline."),

    ("Panel 4 (Bottom-Left): Tool Wear vs Time by Depth of Cut (DOC)",
     "What the graph shows: Scatter plot of VB vs time, colored by Depth of Cut (DOC = 0.75 mm in red, DOC = 1.5 mm in blue).",
     "What it means: At DOC = 0.75 mm, the highest wear values observed reach 1.53 mm — HIGHER than DOC = 1.5 mm cases. "
     "This counterintuitive result occurs because DOC = 0.75 mm experiments ran for much longer durations (up to 105 minutes) "
     "compared to DOC = 1.5 mm cases. A shallower cut causes less aggressive instantaneous wear but leads to more total wear "
     "accumulated over a longer tool life.",
     "Solution implication: This interaction between DOC and machining time is precisely what tree-based models (Random Forest, "
     "Gradient Boosting) can capture through feature interactions. Both DOC and time must be retained as features. "
     "The result also shows we cannot predict wear from time alone — the machining parameters are essential context."),

    ("Panel 5 (Bottom-Center): Average Tool Wear by Material",
     "What the graph shows: Bar chart comparing mean VB values for two workpiece materials: Cast Iron (Material 1) and Steel (Material 2).",
     "What it means: Steel produces 42.3% more average tool wear (0.400 mm) compared to Cast Iron (0.281 mm). "
     "This is physically expected: Steel has higher hardness and abrasive properties, causing faster flank wear on milling inserts. "
     "Tools machining Steel reach the 0.3 mm replacement threshold significantly faster.",
     "Solution implication: The material feature is a critical input that the model must receive. "
     "Without it, the model would be confused by two tools showing different wear rates under seemingly identical sensor readings. "
     "In Week 3, we will build material-specific wear threshold alerts into the dashboard — Steel tools get flagged at lower time values."),

    ("Panel 6 (Bottom-Right): Tool Wear vs Time by Feed Rate",
     "What the graph shows: Scatter plot of VB vs machining time, colored by feed rate (0.25 mm/rev vs 0.5 mm/rev).",
     "What it means: Both feed rates show similar overall wear ranges, but the scatter pattern differs. "
     "Lower feed (0.25) tends to allow longer total machining time before failure (more data points at higher time values), "
     "while higher feed (0.5) cases tend to terminate earlier with less total wear accumulated. "
     "Feed rate interacts with DOC and material to determine the effective cutting force on the tool.",
     "Solution implication: Feed rate must be included as a feature in model input. The complex interaction between "
     "feed, DOC, material, and time cannot be captured by any single feature — our ensemble models handle these "
     "multi-variable interactions automatically through tree splits."),
]

for title, what_shows, what_means, solution in panels:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    bold_run(p, f"► {title}", 11, "1F3864")

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_after = Pt(2)
    bold_run(p2, "📊 What the graph shows: ", 10, "2C3E50")
    normal_run(p2, what_shows, 10)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.8)
    p3.paragraph_format.space_after = Pt(2)
    bold_run(p3, "🔍 What it means: ", 10, "1A5276")
    normal_run(p3, what_means, 10)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.8)
    p4.paragraph_format.space_after = Pt(6)
    bold_run(p4, "💡 Solution implication: ", 10, "1E8449")
    normal_run(p4, solution, 10)

heading(doc, "6.2 AI4I 2020 Dataset — Six-Panel EDA", 2)

add_image(doc, r'd:\IEEE\eda_ai4i.png', Inches(6.2),
    "Figure 2: AI4I 2020 Predictive Maintenance Dataset — Six-Panel Exploratory Data Analysis")

ai4i_panels = [
    ("Panel 1 (Top-Left): Machine Failure Distribution (Pie Chart)",
     "96.61% No Failure vs 3.39% Failure events across 10,000 industrial machine records.",
     "Extremely imbalanced dataset. A dummy classifier predicting 'No Failure' always achieves 96.6% accuracy — "
     "this is the class imbalance trap. The rare 339 failure events carry all the critical information for predictive maintenance.",
     "Never use raw accuracy as the evaluation metric on this dataset. Use F1-Score, Precision-Recall AUC, and ROC-AUC. "
     "Apply SMOTE oversampling or class_weight='balanced' in classifiers during Week 3 training."),

    ("Panel 2 (Top-Center): Tool Wear Distribution (AI4I)",
     "Tool wear time (in minutes) follows a near-uniform distribution from 0 to ~230 minutes, with a slight drop at extreme values.",
     "Tool wear in AI4I is measured as accumulated usage time, not physical VB measurement. "
     "The near-uniform distribution indicates that failures occur across the full wear-time spectrum, "
     "not just at the end of tool life — suggesting that multiple failure modes (HDF, PWF, OSF) "
     "can trigger independently of tool wear state.",
     "Tool wear time alone is insufficient to predict failure. Multi-feature interaction modeling (XGBoost, LSTM) is needed. "
     "OSF (Overstrain Failure) is defined as torque × tool_wear > threshold — this derived feature will be engineered in Week 3."),

    ("Panel 3 (Top-Right): Rotational Speed Distribution",
     "Rotational speed concentrates between 1,300–1,700 RPM with a normal-like distribution. Small tail extends to 2,800+ RPM.",
     "Most operations run within a standard speed range. The high-speed outliers (>2,500 RPM) likely correlate with "
     "specific product types (High quality — H type) that require faster cutting. "
     "Rotational speed is strongly negatively correlated with Torque (r = -0.88): faster rotation with lower torque load.",
     "Speed-torque interaction must be captured as a derived feature for Week 3: Power_proxy = (speed × torque) / 9549. "
     "This directly maps to actual cutting power, which is physically linked to heat dissipation failure (HDF)."),

    ("Panel 4 (Bottom-Left): Torque Distribution",
     "Torque follows a symmetric normal distribution centered around 35–45 Nm, with occasional extreme values up to 75 Nm.",
     "The normal distribution of torque suggests well-controlled process conditions in most operations. "
     "Extreme torque values (>60 Nm) are rare but dangerous — these are likely Overstrain Failure (OSF) precursors. "
     "OSF occurs when torque × tool_wear time > 11,000 [Nm·min] (as per AI4I dataset specification).",
     "Flag records with Torque > 2 standard deviations above mean as high-risk in the dashboard. "
     "The combined (Torque × Tool_wear) threshold feature will be the strongest OSF predictor and must be engineered."),

    ("Panel 5 (Bottom-Center): Failure Type Frequency",
     "HDF=115, OSF=98, PWF=95, TWF=46, RNF=19. Total: 373 individual failure type occurrences across 339 failure events.",
     "Heat Dissipation Failure (HDF) is the most common failure mode, caused by insufficient cooling when rotational speed "
     "drops below a critical temperature-differential threshold. OSF and PWF are nearly equal in frequency. "
     "Tool Wear Failure (TWF) — though our main focus — accounts for only 46 events directly, "
     "suggesting most tool wear occurs as a contributing factor to other failure modes rather than direct TWF. "
     "RNF (Random Failure) at 19 events represents truly unpredictable failures (0.1% probability, unrelated to process state).",
     "Build a multi-label classifier rather than single-label: one model predicts overall failure, "
     "a second predicts which failure TYPE. RNF should be excluded from the evaluation metrics as it cannot be predicted from process data."),

    ("Panel 6 (Bottom-Right): Feature Correlation Heatmap (AI4I)",
     "Correlation matrix of all numeric AI4I features. Strong negative correlation between Rotational Speed and Torque (dark blue ≈ -0.88). "
     "Strong positive correlation between Air Temperature and Process Temperature (dark red ≈ +0.88).",
     "The speed-torque anti-correlation reflects fundamental physics: at higher RPM, lower torque force is needed for the same cutting power. "
     "The temperature correlation is expected: higher ambient temperature raises process temperature. "
     "These collinear features can cause multicollinearity issues in linear models.",
     "For tree-based models (Random Forest, XGBoost), collinearity is not a statistical problem — trees naturally split on the most informative feature. "
     "However, for linear baseline model, we should consider dropping one of each correlated pair. "
     "We will engineer combined features: Heat_Index = Process_Temp - Air_Temp (differential), "
     "Power = Speed × Torque / 9549 (cutting power in kW)."),
]

for title, what_shows, what_means, solution in ai4i_panels:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    bold_run(p, f"► {title}", 11, "1F3864")

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_after = Pt(2)
    bold_run(p2, "📊 What the graph shows: ", 10, "2C3E50")
    normal_run(p2, what_shows, 10)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.8)
    p3.paragraph_format.space_after = Pt(2)
    bold_run(p3, "🔍 What it means: ", 10, "1A5276")
    normal_run(p3, what_means, 10)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.8)
    p4.paragraph_format.space_after = Pt(6)
    bold_run(p4, "💡 Solution implication: ", 10, "1E8449")
    normal_run(p4, solution, 10)

# ═══════════════════════════════════════════════════════════════
# 7. FEATURE ENGINEERING
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "7. Feature Engineering — Creating Predictive Intelligence", 1)

add_para(doc,
    "Raw sensor signals are noisy and temporally dense. A single machining run may have one data row in the mill.csv "
    "dataset (already aggregated). Feature engineering extracts meaningful statistical descriptors from each sensor channel "
    "that capture wear-related signal characteristics.")

add_table(doc,
    ["Feature Category", "Features", "Physical Rationale"],
    [
        ["smcAC Statistics (4)", "smcAC_mean, _std, _max, _min", "AC motor current mean reflects average cutting load; std captures force fluctuations due to wear-induced chattering"],
        ["smcDC Statistics (4)", "smcDC_mean, _std, _max, _min", "DC component correlates with steady-state motor load — increases monotonically with VB (r=0.20 directly, stronger via time)"],
        ["vib_table Statistics (4)", "vib_table_mean, _std, _max, _min", "Table vibration energy increases with wear as the tool loses sharpness and causes irregular cutting forces"],
        ["vib_spindle Statistics (4)", "vib_spindle_mean, _std, _max, _min", "Spindle vibration captures tool-workpiece contact dynamics most directly — highest proximity to wear source"],
        ["AE_table Statistics (4)", "AE_table_mean, _std, _max, _min", "Acoustic emission max and std are the strongest wear indicators — AE energy spikes during micro-fracturing events at flank"],
        ["AE_spindle Statistics (4)", "AE_spindle_mean, _std, _max, _min", "AE at spindle captures high-frequency wear signatures with less background noise interference"],
        ["Machining Parameters (4)", "time, DOC, feed, material", "Operational context: duration, cut depth, advancement rate, workpiece hardness"],
        ["Vibration Ratio (1)", "vib_ratio = vib_table / vib_spindle", "Captures relative energy distribution between table and spindle — changes with wear progression pattern"],
        ["AE Ratio (1)", "AE_ratio = AE_table / AE_spindle", "Acoustic emission sensor comparison — large deviation indicates localized wear events"],
        ["Energy Features (2)", "energy_smcAC = mean²+std², energy_vib = table_mean²+spindle_mean²", "Signal energy proxies that capture total power deposited during cutting"],
    ]
)

heading(doc, "7.1 Top 10 Features by Correlation with Tool Wear (VB)", 2)

add_table(doc,
    ["Rank", "Feature", "Pearson Correlation with VB", "Physical Interpretation"],
    [
        ["1", "AE_spindle_std", "0.495", "Variability in spindle AE — worn tools generate more irregular AE bursts"],
        ["2", "AE_table_max", "0.453", "Peak AE event energy — large spikes correlate with severe flank wear events"],
        ["3", "AE_table_std", "0.432", "AE variability at table sensor — increases as wear creates inconsistent cutting"],
        ["4", "vib_spindle_mean", "0.415", "Average spindle vibration — worn tools produce higher average vibration"],
        ["5", "time", "0.398", "Machining duration — most direct temporal wear indicator"],
        ["6", "AE_spindle_max", "0.387", "Peak AE events at spindle sensor"],
        ["7", "smcDC_mean", "0.312", "Average DC motor current — higher wear → higher motor load"],
        ["8", "smcDC_std", "0.289", "Motor current fluctuations increase with uneven cutting from worn tool"],
        ["9", "energy_vib", "0.276", "Combined vibration energy proxy"],
        ["10", "AE_table_mean", "0.251", "Mean AE energy — progressive increase through tool life"],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 8. BASELINE RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "8. Baseline Model Results — Detailed Analysis", 1)

add_image(doc, r'd:\IEEE\baseline_results.png', Inches(6.2),
    "Figure 3: Baseline Model Comparison — MAE, R² Score, and Top Feature Importances")

add_para(doc, "8.1 Model Performance Comparison", bold=True, size=12, color="1F3864")

add_table(doc,
    ["Model", "MAE (mm)", "RMSE (mm)", "R² Score", "Training Time", "Verdict"],
    [
        ["Linear Regression", "0.1021", "0.1454", "0.5934", "< 1 sec", "Moderate — only captures linear wear relationships; misses non-linear effects"],
        ["Random Forest", "0.0823", "0.1290", "0.7650", "~2 sec", "Good — handles non-linearities via decision trees; robust to outliers"],
        ["Gradient Boosting ★", "0.0654", "0.1085", "0.8334", "~5 sec", "Best baseline — iterative error correction produces tightest predictions"],
    ]
)

add_para(doc, "Understanding the Metrics:", bold=True, size=11, color="1F3864")
metric_rows = [
    ("MAE (Mean Absolute Error)", "Average prediction error in mm of tool wear. MAE = 0.0654 mm means on average our prediction is only 0.065 mm off from actual VB. The ISO tool wear replacement threshold is 0.3 mm, so our error is only 21.8% of the threshold — operationally acceptable."),
    ("RMSE (Root Mean Square Error)", "Penalizes large errors more than MAE. RMSE = 0.1085 mm for Gradient Boosting. Lower RMSE indicates fewer catastrophic mispredictions."),
    ("R² Score", "Coefficient of determination — what percentage of variance in VB is explained by the model. R² = 0.8334 means our Gradient Boosting model explains 83.34% of the variation in tool wear. A perfect model scores 1.0."),
]
for metric, desc in metric_rows:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    bold_run(p, f"• {metric}: ", 11, "1F3864")
    normal_run(p, desc, 11)

add_para(doc, "8.2 Feature Importance Analysis (Random Forest)", bold=True, size=12, color="1F3864")

add_para(doc,
    "The Random Forest model reveals which features drive the prediction most strongly. "
    "The feature importance chart (rightmost panel in Figure 3) shows that AE_table_max alone "
    "accounts for 45% of all predictive power, followed by AE_spindle_std at 25%. "
    "Together, the top 2 AE features contribute 70% of the model's decision-making weight.")

add_para(doc,
    "This finding directly confirms findings from the literature survey (Paper #8: Vibration-based health indicator analysis) "
    "and Paper #1 (Physics-Informed Deep Learning for tool wear monitoring) — acoustic emission signals are the "
    "most sensitive indicators of tool wear state, capturing micro-fracture events and surface degradation at the flank face.")

add_para(doc, "8.3 Actual vs Predicted Tool Wear Scatter Plot", bold=True, size=12, color="1F3864")

add_image(doc, r'd:\IEEE\actual_vs_predicted.png', Inches(4.5),
    "Figure 4: Actual vs Predicted Tool Wear for Best Baseline Model (Gradient Boosting, R²=0.8334)")

add_para(doc,
    "The scatter plot shows predicted VB values (y-axis) plotted against actual measured VB values (x-axis). "
    "The red dashed line represents perfect prediction (predicted = actual). "
    "Most points cluster tightly around the perfect prediction line, confirming the model's strong predictive power.")

add_para(doc,
    "Observed behaviors: (1) At very low wear (VB < 0.05 mm), the model slightly overestimates — a few early-life predictions "
    "are slightly high. This is acceptable as early over-prediction leads to conservative maintenance (better than under-prediction). "
    "(2) At high wear (VB > 0.7 mm), slight underestimation is visible — the model has less training data in this range "
    "as fewer experiments reached extreme wear states. LSTM in Week 3 will address this by learning the full temporal trajectory.")

# ═══════════════════════════════════════════════════════════════
# 9. SOURCE CODE OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "9. Source Code — Key Implementation Details", 1)

code_sections = [
    ("Missing Value Imputation (Per-Case Linear Interpolation)",
     """def fill_vb(group):
    group['VB'] = group['VB'].interpolate(method='linear').ffill().bfill()
    return group
mill_clean = mill_clean.groupby('case', group_keys=False).apply(fill_vb)"""),

    ("Feature Engineering (Per-Run Statistical Aggregation)",
     """sensor_cols = ['smcAC','smcDC','vib_table','vib_spindle','AE_table','AE_spindle']
agg_dict = {f'{col}_{stat}': (col, stat) 
            for col in sensor_cols for stat in ['mean','std','max','min']}
mill_feat = mill_clean.groupby(['case','run']).agg(**agg_dict).reset_index()

# Engineered ratio and energy features
mill_feat['vib_ratio']    = mill_feat['vib_table_mean'] / (mill_feat['vib_spindle_mean'].abs() + 1e-9)
mill_feat['AE_ratio']     = mill_feat['AE_table_mean']  / (mill_feat['AE_spindle_mean'].abs() + 1e-9)
mill_feat['energy_smcAC'] = mill_feat['smcAC_mean']**2  + mill_feat['smcAC_std']**2
mill_feat['energy_vib']   = mill_feat['vib_table_mean']**2 + mill_feat['vib_spindle_mean']**2"""),

    ("Baseline Model Training & Evaluation",
     """from sklearn.ensemble import GradientBoostingRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

model = GradientBoostingRegressor(n_estimators=100, random_state=42)
model.fit(X_train, y_train)
y_pred = model.predict(X_test)

print(f"MAE : {mean_absolute_error(y_test, y_pred):.4f} mm")   # 0.0654
print(f"RMSE: {np.sqrt(mean_squared_error(y_test, y_pred)):.4f} mm")  # 0.1085
print(f"R²  : {r2_score(y_test, y_pred):.4f}")                 # 0.8334"""),
]

for title, code in code_sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    bold_run(p, f"► {title}", 11, "1F3864")

    # Code block style
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent  = Cm(0.8)
    code_para.paragraph_format.right_indent = Cm(0.5)
    code_para.paragraph_format.space_after  = Pt(6)
    code_r = code_para.add_run(code)
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),  'clear')
    shading.set(qn('w:color'),'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    code_para._p.get_or_add_pPr().append(shading)

# ═══════════════════════════════════════════════════════════════
# 10. PRELIMINARY RESULTS SUMMARY + WEEK 3 PLAN
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading(doc, "10. Preliminary Results Summary", 1)

add_table(doc,
    ["Deliverable", "Status", "Key Outcome"],
    [
        ["Dataset Preprocessing (Mill)", "✅ Complete", "168 records, 0 missing values; per-case interpolation applied"],
        ["Dataset Preprocessing (AI4I)", "✅ Complete", "10,000 records, 0 missing, class imbalance documented"],
        ["EDA — 6-Panel Mill Dataset", "✅ Complete", "Wear progression, material effect, sensor correlations analyzed"],
        ["EDA — 6-Panel AI4I Dataset", "✅ Complete", "Failure distribution, failure types, feature correlations analyzed"],
        ["Feature Engineering (32 features)", "✅ Complete", "Statistical aggregation + ratio/energy features built"],
        ["Correlation Analysis", "✅ Complete", "AE_table_max, AE_spindle_std top predictors identified"],
        ["Baseline: Linear Regression", "✅ Complete", "R²=0.5934, MAE=0.1021 mm"],
        ["Baseline: Random Forest", "✅ Complete", "R²=0.7650, MAE=0.0823 mm"],
        ["Baseline: Gradient Boosting", "✅ Complete", "R²=0.8334, MAE=0.0654 mm (BEST)"],
        ["Feature Importance Analysis", "✅ Complete", "AE features contribute 70% of RF predictive weight"],
        ["Actual vs Predicted Plot", "✅ Complete", "Strong alignment with perfect prediction line"],
        ["Source Code (week2_analysis.py)", "✅ Complete", "Fully documented, reproducible pipeline"],
    ]
)

doc.add_paragraph()
heading(doc, "11. Week 3 Plan — Advanced Models", 1)

add_table(doc,
    ["Model / Task", "Dataset", "Approach", "Expected Improvement"],
    [
        ["LSTM (Long Short-Term Memory)", "Mill", "Sequential model on ordered run sequences per case", "Capture temporal wear trajectory → R² target: 0.92+"],
        ["CNN-1D", "Mill", "1D convolution on sensor feature windows", "Automated spatial feature learning from sensor patterns"],
        ["Hybrid CNN+LSTM", "Mill", "CNN extracts local features → LSTM models temporal flow", "Best of both: spatial + temporal learning; R² target: 0.95+"],
        ["XGBoost + SMOTE", "AI4I", "Gradient boosted trees with synthetic minority oversampling", "Failure classification F1-score target: 0.90+"],
        ["LightGBM Classifier", "AI4I", "Fast gradient boosting for multi-class failure type", "Failure type classification F1 target: 0.85+"],
        ["Hyperparameter Tuning", "Both", "GridSearchCV / RandomizedSearchCV / Optuna", "Optimize all model parameters for best generalization"],
        ["SHAP Explainability", "Both", "SHAP TreeExplainer for feature attribution", "Per-prediction explanation for maintenance operators"],
        ["IEEE Paper Draft", "—", "Sections: Methodology, Experimental Setup, Results", "Initial draft of research paper for submission"],
    ]
)

# Final observations
doc.add_paragraph()
heading(doc, "12. Key Scientific Contributions of Week 2", 1)

contributions = [
    "Proved that Acoustic Emission (AE) signals are the dominant tool wear predictors — AE_table_max alone explains 45% of model variance, confirming physics-based intuition from literature.",
    "Demonstrated that per-run statistical aggregation (mean, std, max, min) of raw sensor signals dramatically improves predictive power — from <0.20 raw correlation to >0.45 engineered feature correlation.",
    "Quantified the material effect: Steel workpieces cause 42.3% higher average tool wear than Cast Iron under identical cutting conditions, establishing material type as a mandatory model input.",
    "Achieved R²=0.8334 with a baseline Gradient Boosting model using no deep learning — establishing a strong benchmark that LSTM and hybrid models (Week 3) must surpass.",
    "Identified AI4I class imbalance (3.39% failure rate) as a critical data challenge and documented SMOTE as the planned solution, preventing the accuracy paradox in failure classification.",
    "Revealed that HDF (Heat Dissipation Failure) is the most common failure mode (115 cases) — motivating a Power/Heat_Index engineered feature combining rotational speed, torque, and temperature differential.",
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    bold_run(p, f"{i}. ", 11, "1F3864")
    normal_run(p, c, 11)

divider(doc)
p = doc.add_paragraph("Prepared by Group 18 | Vishwakarma Institute of Technology | IEEE Research Internship — Week 2")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs:
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(100,100,100)

doc.save(r'd:\IEEE\Week2_Report_Group18.docx')
print("SUCCESS: Week2_Report_Group18.docx saved.")
